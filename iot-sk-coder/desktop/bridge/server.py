"""
IOT-ST-KITTS-CODE desktop bridge server.

Runs as a local HTTP server on 127.0.0.1:7823.
Electron's renderer communicates with this via fetch/EventSource.

Routes
──────
GET  /health              - alive check, returns version + brain stats
GET  /brain               - full brain JSON
GET  /config              - current model config
POST /config              - save model config
POST /auth                - verify team token (body: {token})
POST /model/validate      - validate API key (body: model_cfg)
POST /chat                - SSE stream: LLM response + tool use loop
POST /run-brief           - extract agent task brief from conversation
POST /brain/distil        - distil session into brain
GET  /sessions            - list recent sessions with preview
GET  /session?file=...    - load a specific session
GET  /sessions/search?q=  - full-text search across all sessions
GET  /git-context?cwd=... - git status/diff/log for the workspace
"""
from __future__ import annotations

import base64 as _base64
import csv as _csv
import importlib.util
import io
import json
import math
import os
import queue
import re
import subprocess
import sys
import threading
import time
import uuid
import difflib
from collections import defaultdict
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path

# ── Bootstrap path so we can import istkc ────────────────────────────────────
_HERE   = Path(__file__).resolve().parent   # .../bridge/
_IS_WIN = sys.platform == "win32"

# Packaged app:  extraResources puts istkc/ next to bridge/ inside resources/
# Dev layout:    bridge/ is desktop/bridge/; istkc/ is iot-sk-coder/istkc/
#                so we must go up two levels (desktop → iot-sk-coder).
_PARENT = _HERE.parent                      # resources/  OR  desktop/
_ROOT   = _PARENT if (_PARENT / "istkc").exists() else _PARENT.parent
sys.path.insert(0, str(_ROOT))

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

from istkc.brain import Brain
from istkc.auth  import _sha256, _expected_token_hash, _save_auth, is_authenticated, AUTH_FILE, DATA_DIR

_brain = Brain()
_lock  = threading.Lock()              # serialise brain writes
_tl    = threading.local()             # per-thread: sse_cb, model_cfg

PORT = 7823

# App version that spawned this bridge (upgrade handshake — Electron replaces
# any running bridge whose app_version doesn't match its own).
_APP_VERSION = os.environ.get("ISTKC_APP_VERSION", "")


# ── Core dependency bootstrap ─────────────────────────────────────────────────
# Fresh machines (especially macOS) won't have the 'openai' package in the
# system Python, and the packaged app can't ship it inside resources. Install
# it on demand so first-run "No module named 'openai'" never reaches the user.
_ensure_lock = threading.Lock()


def _ensure_openai() -> str | None:
    """Make sure 'openai' is importable, pip-installing it on first use.
    Returns None on success, or an actionable error message."""
    try:
        import openai  # noqa: F401
        return None
    except ImportError:
        pass
    with _ensure_lock:                       # one install at a time
        try:
            import openai  # noqa: F401     # another thread may have finished it
            return None
        except ImportError:
            pass
        print("[bridge] 'openai' package missing — installing…", flush=True)
        # Try each pip mode until one works:
        #   --user                      → classic per-user site-packages
        #   (bare)                      → virtualenvs, where --user is invalid
        #   --break-system-packages    → PEP 668 "externally-managed" Pythons
        #                                 (Homebrew/Debian) that refuse the above
        _variants = (
            ["--user"],
            [],
            ["--user", "--break-system-packages"],
            ["--break-system-packages"],
        )
        last_err = ""
        for extra in _variants:
            try:
                r = subprocess.run(
                    [sys.executable, "-m", "pip", "install", "--quiet", *extra, "openai>=1.0"],
                    capture_output=True, text=True, timeout=240,
                )
                if r.returncode == 0:
                    break
                err_txt = (r.stderr or r.stdout or "").strip()
                # pip < 23.0 rejects --break-system-packages with a usage error;
                # never let that overwrite the REAL failure (network/permissions)
                # captured from an earlier variant.
                _usage_noise = "no such option" in err_txt.lower()
                if err_txt and (not _usage_noise or not last_err):
                    last_err = err_txt[-400:]
            except Exception as e:
                if not last_err:
                    last_err = str(e)
                continue
        importlib.invalidate_caches()
        try:
            import openai  # noqa: F401
            print("[bridge] 'openai' installed OK", flush=True)
            return None
        except ImportError:
            detail = f"\n\npip said:\n{last_err}" if last_err else ""
            return (
                "The Python package 'openai' is not installed and automatic install "
                "failed (no internet, or pip blocked?). Open a terminal and run:\n"
                f"    {sys.executable} -m pip install openai\n"
                "then try again." + detail
            )


# ── CORS + JSON helpers ───────────────────────────────────────────────────────

def _cors_headers(handler: BaseHTTPRequestHandler) -> None:
    handler.send_header("Access-Control-Allow-Origin",  "*")
    handler.send_header("Access-Control-Allow-Methods", "GET, POST, OPTIONS")
    handler.send_header("Access-Control-Allow-Headers", "Content-Type")


def _send_json(handler: BaseHTTPRequestHandler, data: dict, status: int = 200) -> None:
    body = json.dumps(data, ensure_ascii=False).encode("utf-8")
    handler.send_response(status)
    _cors_headers(handler)
    handler.send_header("Content-Type",   "application/json; charset=utf-8")
    handler.send_header("Content-Length", str(len(body)))
    handler.end_headers()
    handler.wfile.write(body)


def _read_body(handler: BaseHTTPRequestHandler) -> dict:
    length = int(handler.headers.get("Content-Length", 0))
    return json.loads(handler.rfile.read(length)) if length else {}


def _config_path() -> Path:
    return DATA_DIR / "config.json"


def _load_cfg() -> dict:
    p = _config_path()
    if p.exists():
        try:
            return json.loads(p.read_text("utf-8"))
        except Exception:
            pass
    return {}


def _save_cfg(cfg: dict) -> None:
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    _config_path().write_text(json.dumps(cfg, indent=2), encoding="utf-8")


# ── Dynamic skill loader ──────────────────────────────────────────────────────

SKILLS_DIR = _HERE / "skills"
SKILLS_DIR.mkdir(exist_ok=True)

_skills: dict[str, dict] = {}   # tool-name → {definition, execute}


def _load_skills() -> int:
    """Scan SKILLS_DIR, import every *.py that exposes TOOL_DEFINITION + execute()."""
    global _skills
    loaded: dict[str, dict] = {}
    for f in sorted(SKILLS_DIR.glob("*.py")):
        if f.name.startswith("_"):
            continue
        try:
            spec = importlib.util.spec_from_file_location(f"nova_skill_{f.stem}", f)
            mod  = importlib.util.module_from_spec(spec)
            spec.loader.exec_module(mod)                        # type: ignore[union-attr]
            if hasattr(mod, "TOOL_DEFINITION") and hasattr(mod, "execute"):
                n = mod.TOOL_DEFINITION["function"]["name"]
                loaded[n] = {"definition": mod.TOOL_DEFINITION, "execute": mod.execute}
                print(f"[skill] loaded: {n}", flush=True)
        except Exception as exc:
            print(f"[skill] error loading {f.name}: {exc}", flush=True)
    _skills = loaded
    return len(_skills)


def _get_tools() -> list:
    """Return base tools + any loaded skills."""
    return _BASE_TOOLS + [s["definition"] for s in _skills.values()]


# ── Browser control (Playwright) ─────────────────────────────────────────────

_SCREENSHOTS_DIR = _HERE / "screenshots"
_pw_inst         = None   # playwright context
_pw_browser      = None   # Browser instance
_pw_page         = None   # current Page
_PW_LOCK         = threading.Lock()


def _ensure_browser():
    """Return a live Playwright Page, launching Chromium if needed."""
    global _pw_inst, _pw_browser, _pw_page
    try:
        from playwright.sync_api import sync_playwright
    except ImportError:
        raise RuntimeError(
            "playwright not installed — run: pip install playwright && playwright install chromium"
        )
    with _PW_LOCK:
        # Check if existing page is still usable
        alive = False
        if _pw_page is not None and _pw_browser is not None:
            try:
                alive = _pw_browser.is_connected()
            except Exception:
                alive = False

        if not alive:
            # Clean up stale instances
            for obj in (_pw_page, _pw_browser, _pw_inst):
                if obj is not None:
                    try:
                        obj.close() if hasattr(obj, 'close') else obj.stop()
                    except Exception:
                        pass
            _pw_inst    = sync_playwright().start()
            _pw_browser = _pw_inst.chromium.launch(
                headless=False,
                args=["--no-sandbox", "--disable-dev-shm-usage"],
            )
            _pw_page = _pw_browser.new_page(viewport={"width": 1280, "height": 800})

        return _pw_page


def _browser_screenshot() -> str:
    """Take a screenshot of the current page and return a base64 PNG data URL."""
    _SCREENSHOTS_DIR.mkdir(exist_ok=True)
    ts   = int(time.time() * 1000)
    path = _SCREENSHOTS_DIR / f"shot_{ts}.png"
    page = _ensure_browser()
    page.screenshot(path=str(path), full_page=False)
    b64 = _base64.b64encode(path.read_bytes()).decode()
    return "data:image/png;base64," + b64


def _close_browser():
    global _pw_inst, _pw_browser, _pw_page
    with _PW_LOCK:
        for obj, method in [(_pw_page, 'close'), (_pw_browser, 'close'), (_pw_inst, 'stop')]:
            if obj is not None:
                try:
                    getattr(obj, method)()
                except Exception:
                    pass
        _pw_page = _pw_browser = _pw_inst = None


# ── Document creation helpers ─────────────────────────────────────────────────

def _doc_create_docx(path: Path, title: str, blocks: list) -> str:
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        return "python-docx not installed — run: pip install python-docx"
    doc = Document()
    if title:
        doc.add_heading(title, 0)
    for blk in blocks:
        t = blk.get("type", "paragraph")
        if t == "heading":
            doc.add_heading(blk.get("text", ""), level=min(int(blk.get("level", 1)), 9))
        elif t == "paragraph":
            doc.add_paragraph(blk.get("text", ""))
        elif t == "bullet_list":
            for item in blk.get("items", []):
                doc.add_paragraph(str(item), style="List Bullet")
        elif t == "table":
            hdrs = blk.get("headers", [])
            rows = blk.get("rows", [])
            tbl  = doc.add_table(rows=1 + len(rows), cols=max(len(hdrs), 1))
            tbl.style = "Table Grid"
            for ci, h in enumerate(hdrs):
                tbl.rows[0].cells[ci].text = str(h)
            for ri, row in enumerate(rows):
                for ci, val in enumerate(row):
                    if ci < len(tbl.rows[ri + 1].cells):
                        tbl.rows[ri + 1].cells[ci].text = str(val)
        elif t == "page_break":
            doc.add_page_break()
    doc.save(str(path))
    return f"✓ Created {path}"


def _doc_create_xlsx(path: Path, title: str, blocks: list) -> str:
    try:
        import openpyxl
        from openpyxl.styles import Alignment, Font, PatternFill
    except ImportError:
        return "openpyxl not installed — run: pip install openpyxl"
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = (title or "Sheet1")[:31]
    hdr_font  = Font(bold=True, color="FFFFFF")
    hdr_fill  = PatternFill("solid", fgColor="2E75B6")
    row_num   = 1
    for blk in blocks:
        t = blk.get("type", "table")
        if t == "sheet":
            ws = wb.create_sheet(title=(blk.get("name", f"Sheet{len(wb.sheetnames)}")[:31]))
            row_num = 1
        if t in ("table", "sheet"):
            hdrs = blk.get("headers", [])
            rows = blk.get("rows", [])
            if hdrs:
                for ci, h in enumerate(hdrs, 1):
                    cell = ws.cell(row=row_num, column=ci, value=h)
                    cell.font, cell.fill = hdr_font, hdr_fill
                    cell.alignment = Alignment(horizontal="center")
                row_num += 1
            for row in rows:
                for ci, val in enumerate(row, 1):
                    ws.cell(row=row_num, column=ci, value=val)
                row_num += 1
            row_num += 1
        elif t == "heading":
            cell = ws.cell(row=row_num, column=1, value=blk.get("text", ""))
            cell.font = Font(bold=True, size=13)
            row_num += 2
        elif t == "paragraph":
            ws.cell(row=row_num, column=1, value=blk.get("text", ""))
            row_num += 1
    for col in ws.columns:
        ml = max((len(str(c.value or "")) for c in col), default=10)
        ws.column_dimensions[col[0].column_letter].width = min(ml + 4, 60)
    wb.save(str(path))
    return f"✓ Created {path} ({len(wb.sheetnames)} sheet(s))"


def _doc_create_pdf(path: Path, title: str, blocks: list) -> str:
    try:
        from fpdf import FPDF
    except ImportError:
        return "fpdf2 not installed — run: pip install fpdf2"
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    if title:
        pdf.set_font("Helvetica", "B", 18)
        pdf.cell(0, 12, title, new_x="LMARGIN", new_y="NEXT", align="C")
        pdf.ln(4)
    for blk in blocks:
        t = blk.get("type", "paragraph")
        if t == "heading":
            sz = [18, 15, 13, 12][min(int(blk.get("level", 1)) - 1, 3)]
            pdf.set_font("Helvetica", "B", sz)
            pdf.ln(3)
            pdf.cell(0, 8, blk.get("text", ""), new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("Helvetica", size=11)
        elif t == "paragraph":
            pdf.set_font("Helvetica", size=11)
            pdf.multi_cell(0, 6, blk.get("text", ""))
            pdf.ln(2)
        elif t == "bullet_list":
            pdf.set_font("Helvetica", size=11)
            for item in blk.get("items", []):
                pdf.cell(6, 6, "•")
                pdf.multi_cell(0, 6, str(item))
        elif t == "table":
            hdrs = blk.get("headers", [])
            rows = blk.get("rows", [])
            if hdrs:
                cw = min(40, 170 // max(len(hdrs), 1))
                pdf.set_font("Helvetica", "B", 9)
                pdf.set_fill_color(46, 117, 182)
                pdf.set_text_color(255, 255, 255)
                for h in hdrs:
                    pdf.cell(cw, 7, str(h)[:22], border=1, fill=True)
                pdf.ln()
                pdf.set_font("Helvetica", size=9)
                pdf.set_text_color(0)
                for row in rows:
                    for val in row:
                        pdf.cell(cw, 6, str(val)[:22], border=1)
                    pdf.ln()
            pdf.ln(2)
        elif t == "page_break":
            pdf.add_page()
    pdf.output(str(path))
    return f"✓ Created {path}"


def _doc_create_html(path: Path, title: str, blocks: list) -> str:
    def esc(s: str) -> str:
        return str(s).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    lines = [
        "<!DOCTYPE html><html lang='en'><head><meta charset='UTF-8'>",
        f"<title>{esc(title or path.stem)}</title><style>",
        "body{font-family:Arial,sans-serif;max-width:960px;margin:40px auto;padding:0 24px;color:#1a1a1a}",
        "h1{color:#1a5276}h2{color:#1f618d}h3{color:#2874a6}",
        "table{border-collapse:collapse;width:100%;margin:16px 0}",
        "th{background:#2e75b6;color:#fff;padding:8px 12px}",
        "td{border:1px solid #d5d8dc;padding:7px 12px}tr:nth-child(even){background:#f2f3f4}",
        "ul{padding-left:24px}li{margin:4px 0}",
        "</style></head><body>",
    ]
    if title:
        lines.append(f"<h1>{esc(title)}</h1>")
    for blk in blocks:
        t = blk.get("type", "paragraph")
        if t == "heading":
            lvl = max(1, min(6, int(blk.get("level", 1)) + 1))
            lines.append(f"<h{lvl}>{esc(blk.get('text',''))}</h{lvl}>")
        elif t == "paragraph":
            lines.append(f"<p>{esc(blk.get('text',''))}</p>")
        elif t == "bullet_list":
            lines.append("<ul>" + "".join(f"<li>{esc(i)}</li>" for i in blk.get("items", [])) + "</ul>")
        elif t == "table":
            hdrs = blk.get("headers", [])
            rows = blk.get("rows", [])
            lines.append("<table>")
            if hdrs:
                lines.append("<thead><tr>" + "".join(f"<th>{esc(h)}</th>" for h in hdrs) + "</tr></thead>")
            lines.append("<tbody>")
            for row in rows:
                lines.append("<tr>" + "".join(f"<td>{esc(v)}</td>" for v in row) + "</tr>")
            lines.append("</tbody></table>")
        elif t == "page_break":
            lines.append("<hr>")
    lines.append("</body></html>")
    path.write_text("\n".join(lines), encoding="utf-8")
    return f"✓ Created {path}"


def _doc_create_csv(path: Path, blocks: list) -> str:
    buf = io.StringIO()
    writer = _csv.writer(buf)
    for blk in blocks:
        t = blk.get("type", "table")
        if t in ("table", "sheet"):
            hdrs = blk.get("headers", [])
            if hdrs:
                writer.writerow(hdrs)
            for row in blk.get("rows", []):
                writer.writerow(row)
            writer.writerow([])
    path.write_text(buf.getvalue(), encoding="utf-8")
    return f"✓ Created {path}"


def _doc_create_md(path: Path, title: str, blocks: list) -> str:
    lines = [f"# {title}\n"] if title else []
    for blk in blocks:
        t = blk.get("type", "paragraph")
        if t == "heading":
            prefix = "#" * min(int(blk.get("level", 1)) + 1, 6)
            lines.append(f"{prefix} {blk.get('text','')}\n")
        elif t == "paragraph":
            lines.append(f"{blk.get('text','')}\n")
        elif t == "bullet_list":
            lines += [f"- {i}" for i in blk.get("items", [])] + [""]
        elif t == "table":
            hdrs = blk.get("headers", [])
            rows = blk.get("rows", [])
            if hdrs:
                lines.append("| " + " | ".join(str(h) for h in hdrs) + " |")
                lines.append("| " + " | ".join("---" for _ in hdrs) + " |")
                for row in rows:
                    lines.append("| " + " | ".join(str(v) for v in row) + " |")
                lines.append("")
        elif t == "page_break":
            lines.append("---\n")
    path.write_text("\n".join(lines), encoding="utf-8")
    return f"✓ Created {path}"


# ── Tool definitions (OpenAI function-calling format) ─────────────────────────

_BASE_TOOLS = [
    {
        "type": "function",
        "function": {
            "name": "read_file",
            "description": (
                "Read the full content of a file from disk. Use this before editing any file "
                "so you always work from the current state, never from memory."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {
                        "type": "string",
                        "description": "File path — absolute or relative to the workspace root.",
                    }
                },
                "required": ["path"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "list_directory",
            "description": (
                "List the immediate children of a directory (files and sub-dirs). "
                "Use to explore project structure without reading every file."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {
                        "type": "string",
                        "description": "Directory path — absolute or relative to workspace root.",
                    }
                },
                "required": ["path"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "run_command",
            "description": (
                "Execute a shell command and capture stdout + stderr. "
                "Use for running tests, builds, linters, installs, git commands, or inspecting "
                "the environment. Keep commands non-interactive. Default timeout is 120 s; for "
                "long jobs (npm/pip install, full test suites, builds) pass a larger 'timeout' "
                "up to 600 s."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "command": {
                        "type": "string",
                        "description": "Shell command to run.",
                    },
                    "cwd": {
                        "type": "string",
                        "description": "Working directory (defaults to workspace root).",
                    },
                    "timeout": {
                        "type": "integer",
                        "description": "Max seconds to wait before killing the command (default 120, max 600).",
                    },
                },
                "required": ["command"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "write_file",
            "description": (
                "Write (or overwrite) a file on disk with the provided content. "
                "Use this to create new files or apply full rewrites. "
                "Creates any missing parent directories automatically."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {
                        "type": "string",
                        "description": "File path — absolute or relative to the workspace root.",
                    },
                    "content": {
                        "type": "string",
                        "description": "Full content to write to the file.",
                    },
                },
                "required": ["path", "content"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "search_code",
            "description": (
                "Search for a regex pattern across all files in the workspace. "
                "Returns matching lines with file paths and line numbers. "
                "Use to find where a symbol is defined or used."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "pattern": {
                        "type": "string",
                        "description": "Grep-style regex pattern to search for.",
                    },
                    "file_pattern": {
                        "type": "string",
                        "description": "Optional glob to filter files, e.g. '*.py' or '*.{js,ts}'.",
                    },
                },
                "required": ["pattern"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "create_document",
            "description": (
                "Create a formatted document: Word (.docx), Excel (.xlsx), PDF, HTML, CSV, or Markdown. "
                "Far more powerful than write_file for structured output with headings, tables, and styling. "
                "Format is auto-detected from the file extension."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {
                        "type": "string",
                        "description": "Output path including extension, e.g. 'report.docx' or 'data.xlsx'.",
                    },
                    "title": {
                        "type": "string",
                        "description": "Document title (used as heading / metadata).",
                    },
                    "content": {
                        "type": "array",
                        "description": (
                            "Array of content blocks. Each block has a 'type' key:\n"
                            "  heading  → {type, level (1-4), text}\n"
                            "  paragraph→ {type, text}\n"
                            "  table    → {type, headers: [...], rows: [[...], ...]}\n"
                            "  bullet_list → {type, items: [...]}\n"
                            "  sheet    → {type, name, headers, rows}  (xlsx only — creates new sheet)\n"
                            "  page_break → {type}"
                        ),
                        "items": {"type": "object"},
                    },
                },
                "required": ["path", "content"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "reload_skills",
            "description": (
                "Hot-reload all skill files from the skills/ directory. "
                "Call this after writing a new skill file to make it available as a tool immediately."
            ),
            "parameters": {"type": "object", "properties": {}, "required": []},
        },
    },
    # ── Browser tools ─────────────────────────────────────────────────────────
    {
        "type": "function",
        "function": {
            "name": "browser_open",
            "description": (
                "Navigate the built-in Chromium browser to a URL, wait for the page to load, "
                "take a screenshot, and return the page title + screenshot for visual verification. "
                "Use this to verify a web app you just built looks and works correctly. "
                "The screenshot appears inline in the chat and the in-app preview pane auto-opens."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "url":  {"type": "string",  "description": "URL to open, e.g. http://localhost:5000"},
                    "wait": {"type": "integer", "description": "Extra ms to wait after load for JS/animations (default 1000)"},
                },
                "required": ["url"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "browser_screenshot",
            "description": "Take a screenshot of the current browser page without navigating. Use to re-check the page after an interaction.",
            "parameters": {"type": "object", "properties": {}, "required": []},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "browser_read",
            "description": "Return the visible text content of the current browser page. Use to verify dynamic content or read form outputs.",
            "parameters": {"type": "object", "properties": {}, "required": []},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "browser_click",
            "description": "Click an element on the current browser page. Provide a CSS selector or text='...' to match by visible text.",
            "parameters": {
                "type": "object",
                "properties": {
                    "selector": {"type": "string", "description": "CSS selector (e.g. '#btn') or text='Submit'"},
                },
                "required": ["selector"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "browser_type",
            "description": "Clear an input field and type text into it on the current browser page.",
            "parameters": {
                "type": "object",
                "properties": {
                    "selector": {"type": "string", "description": "CSS selector of the input/textarea"},
                    "text":     {"type": "string", "description": "Text to type"},
                },
                "required": ["selector", "text"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "browser_close",
            "description": "Close the browser when done testing. Call this to free resources after verification is complete.",
            "parameters": {"type": "object", "properties": {}, "required": []},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "edit_file",
            "description": (
                "Make a surgical text replacement in a file. "
                "Finds exactly one occurrence of old_string and replaces it with new_string. "
                "Fails with an error if old_string is not found or appears more than once, "
                "so always include enough surrounding context to make the match unique."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "path":       {"type": "string", "description": "File path (absolute or relative to workspace)"},
                    "old_string": {"type": "string", "description": "Exact text to find (must match exactly once)"},
                    "new_string": {"type": "string", "description": "Text to replace it with"},
                },
                "required": ["path", "old_string", "new_string"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "web_search",
            "description": (
                "Search the web for documentation, answers, or current information. "
                "Returns a list of results with titles, URLs, and snippets. "
                "Use for finding library docs, error explanations, or any up-to-date information."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "query":       {"type": "string",  "description": "Search query"},
                    "max_results": {"type": "integer", "description": "Maximum number of results to return (default 5)"},
                },
                "required": ["query"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "web_fetch",
            "description": (
                "Fetch a URL and return its content as plain text. "
                "Strips HTML tags so you get readable text. Ideal for reading documentation pages, "
                "README files, or any web content. Respects a 15-second timeout."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "url":      {"type": "string",  "description": "Full URL to fetch (must start with http:// or https://)"},
                    "max_chars": {"type": "integer", "description": "Truncate response to this many characters (default 8000)"},
                },
                "required": ["url"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "web_browse",
            "description": (
                "Research a topic on the web and get a synthesised answer powered by Gemma 4. "
                "Searches the web, fetches the top pages, then uses Google Gemma to read and "
                "summarise the content into a concise, accurate answer. "
                "Use this when you need current information, library docs, error explanations, "
                "or any factual research — it is faster and more accurate than calling "
                "web_search + web_fetch separately."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "query":    {"type": "string",  "description": "Research question or topic to look up"},
                    "focus":    {"type": "string",  "description": "Optional: narrow what to extract (e.g. 'installation steps', 'API usage', 'version compatibility')"},
                },
                "required": ["query"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "git_op",
            "description": (
                "Run a git operation in the workspace. Covers the full workflow: "
                "check status, view diffs, browse history, stage files, commit, push, and pull."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "operation": {
                        "type": "string",
                        "enum": ["status", "diff", "log", "add", "commit", "push", "pull", "branch", "stash"],
                        "description": "Git operation to perform",
                    },
                    "cwd":     {"type": "string", "description": "Working directory (defaults to workspace)"},
                    "message": {"type": "string", "description": "Commit message (required for 'commit')"},
                    "files":   {"type": "array",  "items": {"type": "string"},
                                "description": "Files to stage (for 'add'; omit to stage all changes)"},
                },
                "required": ["operation"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "spawn_agent",
            "description": (
                "Spawn an autonomous background sub-agent to work on a long-running task in parallel. "
                "The agent has the same tools and runs its own agentic loop until complete. "
                "Returns immediately with an agent_id — use agent_status / agent_result to track it."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "task": {"type": "string", "description": "Clear description of what the agent should accomplish"},
                },
                "required": ["task"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "agent_status",
            "description": "Check the status and recent activity log of a spawned background agent.",
            "parameters": {
                "type": "object",
                "properties": {
                    "agent_id": {"type": "string", "description": "The agent ID returned by spawn_agent"},
                },
                "required": ["agent_id"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "agent_result",
            "description": "Get the final result of a completed background agent. Returns an error message if still running.",
            "parameters": {
                "type": "object",
                "properties": {
                    "agent_id": {"type": "string", "description": "The agent ID returned by spawn_agent"},
                },
                "required": ["agent_id"],
            },
        },
    },
]


# ── Tool executor ─────────────────────────────────────────────────────────────

def _py_search(pattern: str, file_pattern: str, workspace: str) -> str:
    """Pure-Python recursive regex search — works on any OS, no grep needed."""
    import fnmatch
    try:
        rx = re.compile(pattern)
    except re.error as e:
        return f"Invalid regex: {e}"
    ws      = Path(workspace)
    matches: list[str] = []
    total   = 0
    try:
        for p in ws.rglob("*"):
            if len(matches) >= 400 or total > 8_000:
                break
            if not p.is_file():
                continue
            if any(d in p.parts for d in _INDEX_SKIP_DIRS):
                continue
            if file_pattern and not fnmatch.fnmatch(p.name, file_pattern):
                continue
            try:
                if p.stat().st_size > 1_000_000:
                    continue
                with p.open("r", encoding="utf-8", errors="replace") as fh:
                    for ln, line in enumerate(fh, 1):
                        if rx.search(line):
                            rel  = p.relative_to(ws)
                            item = f"{rel}:{ln}:{line.rstrip()[:300]}"
                            matches.append(item)
                            total += len(item)
                            if len(matches) >= 400 or total > 8_000:
                                break
            except Exception:
                continue
    except Exception as e:
        return f"Error searching: {e}"
    if not matches:
        return "(no matches)"
    out = "\n".join(matches)
    if total > 8_000:
        out += "\n[...results truncated]"
    return out


def _strip_image_parts(messages: list) -> list:
    """Collapse multimodal image parts to text — for endpoints without vision."""
    out = []
    for m in messages:
        c = m.get("content") if isinstance(m, dict) else None
        if isinstance(c, list):
            texts = [p.get("text", "") for p in c
                     if isinstance(p, dict) and p.get("type") == "text"]
            joined = "\n".join(t for t in texts if t)
            m = {**m, "content": joined or "[image omitted — model has no vision support]"}
        out.append(m)
    return out


def _strip_md_fences(text: str) -> str:
    """Remove a wrapping markdown code fence from output meant to be raw file content."""
    t = text.strip()
    if not t.startswith("```"):
        return text
    first_nl = t.find("\n")
    if first_nl == -1:
        return text
    t = t[first_nl + 1:]
    if t.rstrip().endswith("```"):
        t = t.rstrip()[:-3]
    return t


def _resolve(p: str, workspace: str) -> Path:
    """Resolve p relative to workspace if not absolute."""
    path = Path(p)
    return path if path.is_absolute() else Path(workspace) / path


def _execute_tool(name: str, args: dict, workspace: str) -> str:
    """Run a tool and return the string result (shown to the model as tool output)."""

    if name == "read_file":
        path = _resolve(args.get("path", ""), workspace)
        try:
            content = path.read_text("utf-8", errors="replace")
            if len(content) > 20_000:
                content = content[:20_000] + "\n\n[...file truncated at 20 K chars]"
            ext = path.suffix.lstrip(".") or "text"
            return f"```{ext}\n{content}\n```"
        except FileNotFoundError:
            return f"File not found: {path}"
        except Exception as e:
            return f"Error reading {path}: {e}"

    elif name == "list_directory":
        path = _resolve(args.get("path", workspace), workspace)
        try:
            entries = sorted(path.iterdir(), key=lambda p: (p.is_file(), p.name.lower()))
            lines = []
            for item in entries:
                if item.name.startswith(".") and item.name not in {".env.example", ".gitignore"}:
                    continue
                icon = "📁" if item.is_dir() else "📄"
                lines.append(f"{icon} {item.name}")
            return "\n".join(lines) if lines else "(empty directory)"
        except FileNotFoundError:
            return f"Directory not found: {path}"
        except Exception as e:
            return f"Error listing {path}: {e}"

    elif name == "write_file":
        path = _resolve(args.get("path", ""), workspace)
        content = args.get("content", "")
        try:
            path.parent.mkdir(parents=True, exist_ok=True)
            path.write_text(content, encoding="utf-8")
            return f"✓ Written {len(content)} chars to {path}"
        except Exception as e:
            return f"Error writing {path}: {e}"

    elif name == "run_command":
        cmd = args.get("command", "")
        cwd_arg = args.get("cwd", workspace)
        cwd = str(_resolve(cwd_arg, workspace))
        # Configurable timeout so installs/builds/test-suites don't get cut off
        # at 30 s. Default 120 s, hard ceiling 600 s.
        try:
            _timeout = int(args.get("timeout", 120))
        except (TypeError, ValueError):
            _timeout = 120
        _timeout = max(1, min(_timeout, 600))
        _utf8_env = {**os.environ, "PYTHONUTF8": "1", "PYTHONIOENCODING": "utf-8"}
        try:
            # On Windows, background processes launched via `start /B` or `&`
            # inherit the stdout/stderr pipe handles, causing communicate() to
            # block even after the shell exits.  Workaround: launch via Popen
            # with a dedicated thread to drain pipes, then kill after timeout.
            import threading
            si = None
            cf = 0
            if _IS_WIN:
                import ctypes
                si = subprocess.STARTUPINFO()
                si.dwFlags = subprocess.STARTF_USESHOWWINDOW
                si.wShowWindow = 0  # SW_HIDE
                cf = subprocess.CREATE_NO_WINDOW

            proc = subprocess.Popen(
                cmd, shell=True, cwd=cwd,
                stdout=subprocess.PIPE, stderr=subprocess.PIPE,
                text=True, encoding="utf-8", errors="replace",
                env=_utf8_env,
                startupinfo=si, creationflags=cf,
            )

            stdout_chunks, stderr_chunks = [], []
            _run_sse_cb = getattr(_tl, 'sse_cb', None)
            def _read(pipe, buf, is_stdout=True):
                try:
                    for line in pipe:
                        buf.append(line)
                        if _run_sse_cb and is_stdout:
                            try:
                                _run_sse_cb({"tool_progress": {"name": "run_command", "line": line.rstrip("\n")}})
                            except Exception:
                                pass
                        if sum(len(c) for c in buf) > 10_000:
                            break
                except Exception:
                    pass
                finally:
                    try: pipe.close()
                    except Exception: pass

            t_out = threading.Thread(target=_read, args=(proc.stdout, stdout_chunks, True), daemon=True)
            t_err = threading.Thread(target=_read, args=(proc.stderr, stderr_chunks, False), daemon=True)
            t_out.start(); t_err.start()

            try:
                proc.wait(timeout=_timeout)
            except subprocess.TimeoutExpired:
                # Kill the whole process tree so shells that spawned children
                # (npm, make, pytest -n) don't leave orphans holding the pipes.
                try:
                    if _IS_WIN:
                        subprocess.run(["taskkill", "/PID", str(proc.pid), "/T", "/F"],
                                       capture_output=True)
                    else:
                        proc.kill()
                except Exception:
                    pass
                t_out.join(2); t_err.join(2)
                partial = "".join(stdout_chunks)[-2000:].rstrip()
                return (f"Command timed out after {_timeout} s (killed)."
                        + (f"\n--- partial output ---\n{partial}" if partial else "")
                        + "\nTip: pass a larger 'timeout' (up to 600) for long installs/builds.")

            t_out.join(3); t_err.join(3)
            stdout = "".join(stdout_chunks).rstrip()
            stderr = "".join(stderr_chunks).rstrip()
            out = stdout
            if stderr:
                out += f"\n--- stderr ---\n{stderr}"
            if len(out) > 10_000:
                out = out[:10_000] + "\n[...output truncated]"
            rc = proc.returncode
            return f"exit {rc}\n{out}" if out else f"exit {rc} (no output)"
        except Exception as e:
            return f"Error running command: {e}"

    elif name == "search_code":
        pattern      = args.get("pattern", "")
        file_pattern = args.get("file_pattern", "")
        # Prefer grep where available (fast); fall back to pure-Python search
        # everywhere else — grep does not exist on stock Windows, where this
        # app primarily runs, so the fallback is the main path there.
        if not _IS_WIN:
            try:
                cmd = ["grep", "-rn", "-E", pattern, workspace]
                if file_pattern:
                    cmd = ["grep", "-rn", "-E", "--include", file_pattern, pattern, workspace]
                result = subprocess.run(
                    cmd, capture_output=True, text=True,
                    timeout=15, encoding="utf-8", errors="replace",
                )
                out = result.stdout or "(no matches)"
                if len(out) > 8_000:
                    out = out[:8_000] + "\n[...results truncated]"
                return out
            except subprocess.TimeoutExpired:
                return "Search timed out."
            except FileNotFoundError:
                pass   # no grep — use the Python fallback below
            except Exception as e:
                return f"Error searching: {e}"
        return _py_search(pattern, file_pattern, workspace)

    elif name == "create_document":
        path   = _resolve(args.get("path", ""), workspace)
        title  = args.get("title", path.stem)
        blocks = args.get("content", [])
        fmt    = path.suffix.lower().lstrip(".")
        try:
            path.parent.mkdir(parents=True, exist_ok=True)
            if fmt == "docx":
                return _doc_create_docx(path, title, blocks)
            elif fmt in ("xlsx", "xls"):
                return _doc_create_xlsx(path, title, blocks)
            elif fmt == "pdf":
                return _doc_create_pdf(path, title, blocks)
            elif fmt == "html":
                return _doc_create_html(path, title, blocks)
            elif fmt == "csv":
                return _doc_create_csv(path, blocks)
            elif fmt in ("md", "markdown"):
                return _doc_create_md(path, title, blocks)
            else:
                return f"Unsupported format '{fmt}'. Use: docx, xlsx, pdf, html, csv, md"
        except Exception as e:
            return f"Error creating {fmt}: {e}"

    elif name == "reload_skills":
        count = _load_skills()
        names = list(_skills.keys())
        return (
            f"✓ Loaded {count} skill(s) from {SKILLS_DIR}\n"
            + ("\n".join(f"  • {n}" for n in names) if names else "  (none)")
        )

    # ── Browser tools ─────────────────────────────────────────────────────────
    elif name == "browser_open":
        url      = args.get("url", "")
        wait_ms  = int(args.get("wait", 1000))
        try:
            page = _ensure_browser()
            page.goto(url, wait_until="domcontentloaded", timeout=20000)
            page.wait_for_timeout(wait_ms)
            title    = page.title()
            data_url = _browser_screenshot()
            return (
                f"✓ Navigated to {url}\n"
                f"Title: {title}\n"
                f"PREVIEW_URL:{url}\n"
                f"SCREENSHOT_B64:{data_url}"
            )
        except Exception as e:
            return f"Browser error: {e}"

    elif name == "browser_screenshot":
        try:
            data_url = _browser_screenshot()
            return f"Screenshot taken.\nSCREENSHOT_B64:{data_url}"
        except Exception as e:
            return f"Screenshot error: {e}"

    elif name == "browser_read":
        try:
            page = _ensure_browser()
            text = page.inner_text("body")
            if len(text) > 5000:
                text = text[:5000] + "\n[...truncated at 5 K chars]"
            return f"Page content:\n{text}"
        except Exception as e:
            return f"Error reading page: {e}"

    elif name == "browser_click":
        selector = args.get("selector", "")
        try:
            page = _ensure_browser()
            if selector.startswith("text="):
                page.get_by_text(selector[5:], exact=False).first.click()
            else:
                page.click(selector, timeout=8000)
            page.wait_for_timeout(600)
            data_url = _browser_screenshot()
            return f"✓ Clicked '{selector}'.\nSCREENSHOT_B64:{data_url}"
        except Exception as e:
            return f"Click error on '{selector}': {e}"

    elif name == "browser_type":
        selector = args.get("selector", "")
        text     = args.get("text", "")
        try:
            page = _ensure_browser()
            page.fill(selector, text)
            return f"✓ Typed into '{selector}'"
        except Exception as e:
            return f"Type error on '{selector}': {e}"

    elif name == "browser_close":
        _close_browser()
        return "✓ Browser closed."

    elif name == "edit_file":
        path       = _resolve(args.get("path", ""), workspace)
        old_string = args.get("old_string", "")
        new_string = args.get("new_string", "")
        try:
            content = path.read_text("utf-8", errors="replace")
            count = content.count(old_string)
            if count == 0:
                return f"Error: old_string not found in {path}"
            if count > 1:
                return (
                    f"Error: old_string found {count} times in {path}. "
                    "Add more surrounding context to make the match unique."
                )
            new_content = content.replace(old_string, new_string, 1)
            path.write_text(new_content, encoding="utf-8")
            old_lines = old_string.splitlines(keepends=True)
            new_lines = new_string.splitlines(keepends=True)
            diff_lines = list(difflib.unified_diff(
                old_lines, new_lines,
                fromfile=f"a/{path.name}", tofile=f"b/{path.name}", lineterm="",
            ))
            diff_str = "".join(diff_lines)[:2000]
            return f"✓ Edit applied to {path}\n```diff\n{diff_str}\n```"
        except FileNotFoundError:
            return f"File not found: {path}"
        except Exception as e:
            return f"Error editing {path}: {e}"

    elif name == "web_search":
        query       = args.get("query", "")
        max_results = int(args.get("max_results", 5))
        try:
            # Primary: ddgs library (formerly duckduckgo_search)
            try:
                try:
                    from ddgs import DDGS
                except ImportError:
                    from duckduckgo_search import DDGS
                results = []
                with DDGS() as ddgs:
                    for r in ddgs.text(query, max_results=max_results):
                        results.append(r)
                if results:
                    lines = []
                    for i, r in enumerate(results, 1):
                        title   = r.get("title", "")
                        url     = r.get("href", r.get("url", ""))
                        snippet = r.get("body", r.get("snippet", ""))
                        lines.append(f"{i}. **{title}**\n   {url}\n   {snippet}")
                    return "\n\n".join(lines)
            except Exception:
                pass  # fall through to JSON API fallback
            # Fallback: DuckDuckGo Instant Answer JSON API (no Cloudflare gating)
            import urllib.request, urllib.parse, json as _json
            encoded = urllib.parse.quote_plus(query)
            api_url = f"https://api.duckduckgo.com/?q={encoded}&format=json&no_html=1&skip_disambig=1"
            req = urllib.request.Request(api_url, headers={"User-Agent": "Mozilla/5.0"})
            with urllib.request.urlopen(req, timeout=10) as resp:
                data = _json.loads(resp.read().decode("utf-8", errors="replace"))
            lines = []
            # Abstract result (main answer)
            if data.get("AbstractText"):
                lines.append(f"1. **{data.get('Heading', query)}**\n   {data.get('AbstractURL','')}\n   {data['AbstractText']}")
            # Related topics
            for item in data.get("RelatedTopics", [])[:max_results - len(lines)]:
                if isinstance(item, dict) and item.get("Text"):
                    url = item.get("FirstURL", "")
                    lines.append(f"{len(lines)+1}. **{item['Text'][:80]}**\n   {url}")
            return "\n\n".join(lines) if lines else "No results found. Try web_fetch with a specific docs URL instead."
        except Exception as e:
            return f"Search error: {e}"

    elif name == "web_fetch":
        url       = args.get("url", "")
        max_chars = int(args.get("max_chars", 8000))
        try:
            import urllib.request, re as _re
            req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
            with urllib.request.urlopen(req, timeout=15) as resp:
                content_type = resp.headers.get("Content-Type", "")
                raw = resp.read().decode("utf-8", errors="replace")
            if "html" in content_type or raw.lstrip().startswith("<"):
                # Strip scripts/styles first
                raw = _re.sub(r"(?is)<(script|style)[^>]*>.*?</\1>", " ", raw)
                # Try BeautifulSoup for clean text
                try:
                    from bs4 import BeautifulSoup
                    text = BeautifulSoup(raw, "html.parser").get_text(separator="\n")
                except ImportError:
                    text = _re.sub(r"<[^>]+>", " ", raw)
                # Collapse whitespace
                text = _re.sub(r"[ \t]+", " ", text)
                text = _re.sub(r"\n{3,}", "\n\n", text).strip()
            else:
                text = raw.strip()
            if len(text) > max_chars:
                text = text[:max_chars] + f"\n\n[...truncated at {max_chars} chars]"
            return text
        except Exception as e:
            return f"Fetch error: {e}"

    elif name == "web_browse":
        query = args.get("query", "")
        focus = args.get("focus", "")
        try:
            import urllib.request as _req, json as _json, re as _re

            # ── Step 1: search ──────────────────────────────────────────────
            search_results = []
            try:
                try:
                    from ddgs import DDGS
                except ImportError:
                    from duckduckgo_search import DDGS
                with DDGS() as ddgs:
                    for r in ddgs.text(query, max_results=4):
                        search_results.append(r)
            except Exception:
                pass

            if not search_results:
                return f"web_browse: no search results for '{query}'"

            # ── Step 2: fetch top 2 pages ───────────────────────────────────
            pages = []
            for r in search_results[:2]:
                url = r.get("href", r.get("url", ""))
                if not url:
                    continue
                try:
                    req = _req.Request(url, headers={"User-Agent": "Mozilla/5.0"})
                    with _req.urlopen(req, timeout=12) as resp:
                        raw = resp.read().decode("utf-8", errors="replace")
                    raw = _re.sub(r"(?is)<(script|style)[^>]*>.*?</\1>", " ", raw)
                    try:
                        from bs4 import BeautifulSoup
                        text = BeautifulSoup(raw, "html.parser").get_text(separator="\n")
                    except ImportError:
                        text = _re.sub(r"<[^>]+>", " ", raw)
                    text = _re.sub(r"[ \t]+", " ", text)
                    text = _re.sub(r"\n{3,}", "\n\n", text).strip()
                    pages.append({"url": url, "text": text[:6000]})
                except Exception:
                    continue

            if not pages:
                # Fall back to snippets only
                snippets = "\n\n".join(
                    f"[{r.get('title','')}]\n{r.get('body', r.get('snippet',''))}"
                    for r in search_results[:4]
                )
                pages_text = snippets
            else:
                pages_text = "\n\n---\n\n".join(
                    f"Source: {p['url']}\n\n{p['text']}" for p in pages
                )

            # ── Step 3: synthesise with Gemma 4 ────────────────────────────
            focus_note = f"\nFocus specifically on: {focus}" if focus else ""
            gemma_prompt = (
                f"You are a research assistant. Based on the web content below, "
                f"answer this question concisely and accurately:\n\n"
                f"**Question:** {query}{focus_note}\n\n"
                f"**Web content:**\n{pages_text[:12000]}\n\n"
                f"Give a clear, factual answer. Include code examples if relevant. "
                f"Cite the source URL where helpful."
            )

            cfg = _load_cfg().get("model_cfg", {})
            api_key = cfg.get("api_key", "")

            # Try a free Gemma first, then the user's own model; if synthesis is
            # unavailable, return the raw fetched content instead of an error so
            # the calling model still gets the research material.
            _synth_models = ["google/gemma-3-27b-it:free"]
            if cfg.get("provider") == "openrouter" and cfg.get("api_id"):
                _synth_models.append(cfg["api_id"])

            answer = ""
            for _synth_model in _synth_models:
                try:
                    payload = {
                        "model": _synth_model,
                        "messages": [{"role": "user", "content": gemma_prompt}],
                        "max_tokens": 2048,
                    }
                    req = _req.Request(
                        "https://openrouter.ai/api/v1/chat/completions",
                        data=_json.dumps(payload).encode(),
                        headers={
                            "Authorization": f"Bearer {api_key}",
                            "Content-Type": "application/json",
                            "HTTP-Referer": "http://localhost",
                            "X-Title": "NOVA-web-browse",
                        },
                    )
                    with _req.urlopen(req, timeout=60) as resp:
                        data = _json.loads(resp.read())
                    answer = (data.get("choices") or [{}])[0].get("message", {}).get("content") or ""
                    if answer:
                        break
                except Exception:
                    continue

            sources = "\n".join(f"- {p['url']}" for p in pages)
            if not answer:
                # Synthesis unavailable — hand back the raw content
                answer = f"(synthesis unavailable — raw page content follows)\n\n{pages_text[:6000]}"
            if sources:
                answer = answer.rstrip() + f"\n\n**Sources:**\n{sources}"
            return answer
        except Exception as e:
            return f"web_browse error: {e}"

    # ── git_op ───────────────────────────────────────────────────────────────
    elif name == "git_op":
        operation = args.get("operation", "status")
        cwd = str(_resolve(args.get("cwd", workspace), workspace))
        _READ_OPS = {
            "status":  ["status", "--short"],
            "diff":    ["diff", "--no-color"],
            "log":     ["log", "--oneline", "-20"],
            "branch":  ["branch", "-a"],
            "stash":   ["stash", "list"],
        }
        if operation in _READ_OPS:
            out = _git(_READ_OPS[operation], cwd) or "(nothing)"
            return f"git {operation}:\n{out}"
        elif operation == "add":
            files = args.get("files", ["."])
            if isinstance(files, str):
                files = [files]
            r = subprocess.run(["git", "add"] + files, cwd=cwd,
                               capture_output=True, text=True, timeout=10)
            return f"✓ git add: {r.stdout.strip() or r.stderr.strip() or 'staged'}"
        elif operation == "commit":
            msg = args.get("message", "chore: update")
            r = subprocess.run(["git", "commit", "-m", msg], cwd=cwd,
                               capture_output=True, text=True, timeout=10)
            return r.stdout.strip() or r.stderr.strip() or "Committed."
        elif operation == "push":
            r = subprocess.run(["git", "push"], cwd=cwd,
                               capture_output=True, text=True, timeout=30)
            return r.stdout.strip() or r.stderr.strip() or "Pushed."
        elif operation == "pull":
            r = subprocess.run(["git", "pull"], cwd=cwd,
                               capture_output=True, text=True, timeout=30)
            return r.stdout.strip() or r.stderr.strip() or "Pulled."
        else:
            return f"Unknown git operation '{operation}'. Valid: status, diff, log, add, commit, push, pull, branch, stash"

    # ── spawn_agent ───────────────────────────────────────────────────────────
    elif name == "spawn_agent":
        task_desc = args.get("task", "").strip()
        if not task_desc:
            return "Error: 'task' description is required."
        _model_cfg = getattr(_tl, 'model_cfg', {})
        agent_id   = uuid.uuid4().hex[:8]
        with _agents_lock:
            _agents[agent_id] = {"status": "running", "result": None, "log": []}
        threading.Thread(
            target=_run_agent_bg,
            args=(agent_id, task_desc, _model_cfg, workspace),
            daemon=True,
        ).start()
        return (
            f"✓ Agent spawned — ID: {agent_id}\n"
            f"Task: {task_desc[:120]}\n"
            f"Use agent_status('{agent_id}') to check progress, "
            f"agent_result('{agent_id}') to get the final output."
        )

    # ── agent_status ──────────────────────────────────────────────────────────
    elif name == "agent_status":
        agent_id = args.get("agent_id", "").strip()
        with _agents_lock:
            ag = _agents.get(agent_id)
        if not ag:
            return f"No agent with ID '{agent_id}'. Active agents: {list(_agents.keys())}"
        log_tail = "\n".join(ag.get("log", [])[-6:]) or "(no activity yet)"
        return f"Agent {agent_id}: status={ag['status']}\nRecent activity:\n{log_tail}"

    # ── agent_result ──────────────────────────────────────────────────────────
    elif name == "agent_result":
        agent_id = args.get("agent_id", "").strip()
        with _agents_lock:
            ag = _agents.get(agent_id)
        if not ag:
            return f"No agent with ID '{agent_id}'."
        if ag["status"] == "running":
            return f"Agent {agent_id} is still running. Try agent_status('{agent_id}') for progress."
        return f"Agent {agent_id} [{ag['status']}]:\n{ag.get('result', '(no result)')}"

    # ── Dynamic skills dispatch ───────────────────────────────────────────────
    if name in _skills:
        try:
            return _skills[name]["execute"](args, workspace)
        except Exception as e:
            return f"Skill '{name}' error: {e}"

    return f"Unknown tool: {name}"


# ── Multi-agent registry ──────────────────────────────────────────────────────

_agents: dict[str, dict] = {}
_agents_lock = threading.Lock()


def _run_agent_bg(agent_id: str, task_desc: str, model_cfg: dict, workspace: str) -> None:
    """Background sub-agent: full agentic loop; result stored in _agents."""
    from istkc.models import make_client
    try:
        client  = make_client(model_cfg)
        api_id  = model_cfg.get("api_id", "")
        messages = [
            {"role": "system", "content": (
                "You are a background sub-agent. Complete the given task autonomously using the "
                "available tools. Be thorough and self-sufficient. When finished, provide a clear "
                "summary of everything you accomplished."
            )},
            {"role": "user", "content": task_desc},
        ]
        log: list[str] = []
        full_text = ""

        _tl.model_cfg = model_cfg   # allow nested tool dispatch to see model_cfg

        for _round in range(12):
            resp = client.chat.completions.create(
                model=api_id,
                messages=messages,
                tools=_get_tools(),
                tool_choice="auto",
                max_tokens=4096,
            )
            msg = resp.choices[0].message
            full_text = msg.content or ""
            if full_text:
                log.append(f"[think] {full_text[:180]}")

            if resp.choices[0].finish_reason != "tool_calls" or not msg.tool_calls:
                break

            asst_msg = {
                "role": "assistant",
                "content": full_text or None,
                "tool_calls": [
                    {"id": tc.id, "type": "function",
                     "function": {"name": tc.function.name, "arguments": tc.function.arguments}}
                    for tc in msg.tool_calls
                ],
            }
            tool_results = []
            for tc in msg.tool_calls:
                try:
                    call_args = json.loads(tc.function.arguments or "{}")
                except Exception:
                    call_args = {}
                tool_result = _execute_tool(tc.function.name, call_args, workspace)
                log.append(f"[{tc.function.name}] {tool_result[:120]}")
                tool_results.append({
                    "role": "tool",
                    "tool_call_id": tc.id,
                    "content": tool_result,
                })
            messages = messages + [asst_msg] + tool_results

        with _agents_lock:
            _agents[agent_id]["status"] = "done"
            _agents[agent_id]["result"] = full_text or "Agent completed."
            _agents[agent_id]["log"]    = log

    except Exception as exc:
        with _agents_lock:
            _agents[agent_id]["status"] = "error"
            _agents[agent_id]["result"] = f"Agent error: {exc}"


# ── Git helpers ───────────────────────────────────────────────────────────────

def _git(args: list[str], cwd: str) -> str:
    """Run a git command and return stdout, or empty string on failure."""
    try:
        r = subprocess.run(
            ["git"] + args, cwd=cwd,
            capture_output=True, text=True,
            timeout=8, encoding="utf-8", errors="replace",
        )
        return r.stdout.strip() if r.returncode == 0 else ""
    except Exception:
        return ""


# ── Workspace code indexer (TF-IDF, no ML deps) ──────────────────────────────

_INDEX_EXTENSIONS = {
    '.py', '.js', '.ts', '.jsx', '.tsx', '.mjs', '.cjs',
    '.go', '.rs', '.java', '.c', '.cpp', '.h', '.cs',
    '.rb', '.php', '.swift', '.kt', '.sh', '.bash',
    '.sql', '.html', '.css', '.scss', '.yaml', '.yml',
    '.toml', '.json', '.md', '.txt', '.env.example',
}
_INDEX_SKIP_DIRS = {
    'node_modules', '.git', '__pycache__', 'dist', 'build',
    '.venv', 'venv', 'env', '.next', '.nuxt', 'coverage',
    'target', 'vendor', '.idea', '.vscode',
}


class CodeIndex:
    """Lightweight TF-IDF index over workspace source files."""

    def __init__(self, workspace: str) -> None:
        self.workspace = workspace
        self.chunks: list[dict]        = []
        self.idf:    dict[str, float]  = {}
        self._built  = False
        self._lock   = threading.Lock()

    # ── tokeniser ─────────────────────────────────────────────────────────────
    @staticmethod
    def _tok(text: str) -> list[str]:
        return [t.lower() for t in re.findall(r'[a-zA-Z_][a-zA-Z0-9_]{2,}', text)]

    # ── chunking ──────────────────────────────────────────────────────────────
    def _chunk_file(self, rel_path: str, content: str) -> list[dict]:
        lines   = content.split('\n')
        step    = 40   # lines between chunk starts
        window  = 70   # lines per chunk
        chunks  = []
        for start in range(0, max(1, len(lines)), step):
            end  = min(len(lines), start + window)
            text = '\n'.join(lines[start:end]).strip()
            if text:
                chunks.append({'file': rel_path, 'line': start + 1, 'text': text})
        return chunks

    # ── build ─────────────────────────────────────────────────────────────────
    def build(self) -> dict:
        ws      = Path(self.workspace)
        chunks  = []
        skipped = 0

        for p in ws.rglob('*'):
            if not p.is_file():
                continue
            if any(d in p.parts for d in _INDEX_SKIP_DIRS):
                continue
            if p.suffix not in _INDEX_EXTENSIONS and p.name not in {'.env.example', '.gitignore'}:
                continue
            try:
                sz = p.stat().st_size
                if sz > 400_000:
                    skipped += 1
                    continue
                rel  = str(p.relative_to(ws))
                text = p.read_text('utf-8', errors='replace')
                chunks.extend(self._chunk_file(rel, text))
            except Exception:
                skipped += 1

        # Build IDF
        df = defaultdict(int)
        for ch in chunks:
            for t in set(self._tok(ch['text'])):
                df[t] += 1
        N   = max(1, len(chunks))
        idf = {t: math.log(N / (1 + n)) for t, n in df.items()}

        with self._lock:
            self.chunks = chunks
            self.idf    = idf
            self._built = True

        return {'chunks': len(chunks), 'files': len({c['file'] for c in chunks}), 'skipped': skipped}

    # ── search ────────────────────────────────────────────────────────────────
    def search(self, query: str, top_k: int = 6) -> list[dict]:
        with self._lock:
            if not self._built or not self.chunks:
                return []
            chunks = self.chunks
            idf    = self.idf

        q_toks = self._tok(query)
        if not q_toks:
            return []

        scores: list[tuple[float, int]] = []
        for i, ch in enumerate(chunks):
            ch_toks = self._tok(ch['text'])
            if not ch_toks:
                continue
            tf: dict[str, float] = defaultdict(float)
            for t in ch_toks:
                tf[t] += 1.0 / len(ch_toks)
            score = sum(tf.get(t, 0) * idf.get(t, 0) for t in q_toks)
            if score > 0:
                scores.append((score, i))

        scores.sort(reverse=True)

        results: list[dict] = []
        seen: set[tuple[str, int]] = set()
        for sc, idx in scores[:top_k * 3]:
            ch  = chunks[idx]
            key = (ch['file'], ch['line'] // 60)  # dedup nearby chunks from same file
            if key not in seen:
                seen.add(key)
                results.append({**ch, 'score': round(sc, 4)})
                if len(results) >= top_k:
                    break

        return results


# Global index registry (one per workspace path)
_code_indexes:  dict[str, CodeIndex] = {}
_index_lock     = threading.Lock()


def get_or_create_index(workspace: str) -> CodeIndex:
    with _index_lock:
        if workspace not in _code_indexes:
            _code_indexes[workspace] = CodeIndex(workspace)
        return _code_indexes[workspace]


# ── Background task runner ────────────────────────────────────────────────────

class BgTask:
    """A long-running server-side task that emits SSE progress events."""

    def __init__(self, task_id: str) -> None:
        self.id:    str         = task_id
        self.q:     queue.Queue = queue.Queue()
        self.done:  bool        = False
        self.error: str | None  = None

    def emit(self, **kwargs) -> None:
        self.q.put(kwargs)

    def finish(self, result: str = "") -> None:
        self.done = True
        self.q.put({"done": True, "result": result})

    def fail(self, err: str) -> None:
        self.error = err
        self.done  = True
        self.q.put({"done": True, "error": err})


_bg_tasks:      dict[str, BgTask] = {}
_bg_tasks_lock  = threading.Lock()


def _new_task() -> BgTask:
    tid = str(uuid.uuid4())[:8]
    t   = BgTask(tid)
    with _bg_tasks_lock:
        _bg_tasks[tid] = t
    return t


# ── Request handler ───────────────────────────────────────────────────────────

class Handler(BaseHTTPRequestHandler):

    # Silence access logs in production
    def log_message(self, fmt, *args):  # noqa
        pass

    # ── OPTIONS preflight ─────────────────────────────────────────────────────
    def do_OPTIONS(self):
        self.send_response(204)
        _cors_headers(self)
        self.end_headers()

    # ── GET ───────────────────────────────────────────────────────────────────
    def do_GET(self):
        from urllib.parse import urlparse, parse_qs
        parsed = urlparse(self.path)
        route  = parsed.path
        qs     = parse_qs(parsed.query)

        # ── /health ──────────────────────────────────────────────────────────
        if route == "/health":
            bd = _brain.load()
            _send_json(self, {
                "ok":          True,
                "version":     "0.1.0-istkc",
                "app_version": _APP_VERSION,   # upgrade handshake (see main.js)
                "brain":       _brain.stats(bd),
                "authed":      is_authenticated(),
            })

        # ── /brain ───────────────────────────────────────────────────────────
        elif route == "/brain":
            _send_json(self, _brain.load())

        # ── /codex/status ────────────────────────────────────────────────────
        elif route == "/codex/status":
            from istkc.models import get_codex_status
            _send_json(self, get_codex_status())

        # ── /config ──────────────────────────────────────────────────────────
        elif route == "/config":
            _send_json(self, _load_cfg())

        # ── /sessions ────────────────────────────────────────────────────────
        elif route == "/sessions":
            from istkc.brain import SESSIONS_DIR
            # Sort by last-activity (mtime) so an updated conversation floats to
            # the top — stable-id files are rewritten as the chat continues.
            paths = sorted(
                SESSIONS_DIR.glob("*.jsonl"),
                key=lambda p: p.stat().st_mtime if p.exists() else 0,
                reverse=True,
            )[:40]
            items = []
            for p in paths:
                try:
                    lines   = [l for l in p.read_text("utf-8").splitlines() if l.strip()]
                    msgs    = [json.loads(l) for l in lines]
                    preview = ""
                    for m in msgs:
                        if m.get("role") == "user":
                            c = m.get("content", "")
                            if isinstance(c, list):
                                c = " ".join(
                                    x.get("text", "") for x in c
                                    if isinstance(x, dict) and x.get("type") == "text"
                                )
                            txt = str(c).strip()
                            if txt:
                                preview = txt[:80]
                                break
                    items.append({
                        "file":    p.name,
                        "stem":    p.stem,
                        "preview": preview,
                        "count":   len([m for m in msgs if m.get("role") == "user"]),
                        "mtime":   int(p.stat().st_mtime),
                    })
                except Exception:
                    pass
            _send_json(self, {"sessions": items})

        # ── /sessions/search ─────────────────────────────────────────────────
        elif route == "/sessions/search":
            q = (qs.get("q") or [""])[0].strip().lower()
            from istkc.brain import SESSIONS_DIR
            results = []
            if q:
                for p in sorted(SESSIONS_DIR.glob("*.jsonl"), reverse=True)[:60]:
                    try:
                        raw = p.read_text("utf-8")
                        if q not in raw.lower():
                            continue
                        lines   = [l for l in raw.splitlines() if l.strip()]
                        msgs    = [json.loads(l) for l in lines]
                        preview = ""
                        for m in msgs:
                            c = m.get("content", "")
                            if isinstance(c, list):
                                c = " ".join(
                                    x.get("text", "") for x in c
                                    if isinstance(x, dict) and x.get("type") == "text"
                                )
                            if q in str(c).lower():
                                idx = str(c).lower().find(q)
                                snippet = str(c)[max(0, idx - 25): idx + 80].strip()
                                preview = f"…{snippet}…"
                                break
                        results.append({
                            "file":    p.name,
                            "stem":    p.stem,
                            "preview": preview or p.stem,
                            "count":   len([m for m in msgs if m.get("role") == "user"]),
                            "mtime":   int(p.stat().st_mtime),
                        })
                    except Exception:
                        pass
            _send_json(self, {"results": results[:15]})

        # ── /session ─────────────────────────────────────────────────────────
        elif route == "/session":
            from istkc.brain import SESSIONS_DIR
            # basename only — the filename comes from the client, block ../ traversal
            fname = Path((qs.get("file") or [""])[0]).name
            if fname:
                p = SESSIONS_DIR / fname
                if p.exists():
                    lines = [l for l in p.read_text("utf-8").splitlines() if l.strip()]
                    msgs  = [json.loads(l) for l in lines]
                    _send_json(self, {"messages": msgs})
                else:
                    _send_json(self, {"error": "not found"}, 404)
            else:
                _send_json(self, {"error": "file param required"}, 400)

        # ── /git-context ──────────────────────────────────────────────────────
        elif route == "/git-context":
            cwd = (qs.get("cwd") or [""])[0] or str(Path.home())
            branch    = _git(["branch", "--show-current"], cwd)
            status    = _git(["status", "--short"], cwd)
            diff_stat = _git(["diff", "HEAD", "--stat", "--no-color"], cwd)
            log       = _git(["log", "--oneline", "-7"], cwd)
            stash     = _git(["stash", "list"], cwd)
            _send_json(self, {
                "ok":        bool(branch),
                "branch":    branch,
                "status":    status,
                "diff_stat": diff_stat,
                "log":       log,
                "stash":     stash,
            })

        # ── /workspace/search ─────────────────────────────────────────────────
        elif route == "/workspace/search":
            workspace = (qs.get("workspace") or [""])[0]
            query     = (qs.get("q") or [""])[0].strip()
            top_k     = int((qs.get("k") or ["6"])[0])
            if not workspace or not query:
                _send_json(self, {"results": [], "indexed": False})
            else:
                idx = get_or_create_index(workspace)
                results = idx.search(query, top_k=top_k)
                _send_json(self, {
                    "results": results,
                    "indexed": idx._built,
                    "total_chunks": len(idx.chunks),
                })

        # ── /workspace/index/status ───────────────────────────────────────────
        elif route == "/workspace/index/status":
            workspace = (qs.get("workspace") or [""])[0]
            if not workspace:
                _send_json(self, {"indexed": False, "chunks": 0})
            else:
                idx = get_or_create_index(workspace)
                _send_json(self, {
                    "indexed": idx._built,
                    "chunks":  len(idx.chunks),
                    "files":   len({c['file'] for c in idx.chunks}),
                })

        # ── /task/stream/<id> ─────────────────────────────────────────────────
        elif route.startswith("/task/stream/"):
            tid  = route.split("/")[-1]
            with _bg_tasks_lock:
                task = _bg_tasks.get(tid)
            if not task:
                _send_json(self, {"error": "task not found"}, 404)
                return

            self.send_response(200)
            _cors_headers(self)
            self.send_header("Content-Type",  "text/event-stream; charset=utf-8")
            self.send_header("Cache-Control", "no-cache")
            self.send_header("X-Accel-Buffering", "no")
            self.end_headers()

            def sse(data: dict) -> None:
                line = f"data: {json.dumps(data)}\n\n"
                self.wfile.write(line.encode("utf-8"))
                self.wfile.flush()

            while not task.done:
                try:
                    evt = task.q.get(timeout=0.5)
                    sse(evt)
                    if evt.get("done"):
                        break
                except queue.Empty:
                    sse({"ping": True})  # keep-alive

        # ── /screenshots/<file> ───────────────────────────────────────────────
        elif route.startswith("/screenshots/"):
            # basename only — block ../ traversal from the URL path
            fname = Path(route.split("/screenshots/", 1)[1]).name
            fpath = _SCREENSHOTS_DIR / fname
            if fpath.exists() and fpath.suffix.lower() == ".png":
                data = fpath.read_bytes()
                self.send_response(200)
                _cors_headers(self)
                self.send_header("Content-Type",   "image/png")
                self.send_header("Content-Length", str(len(data)))
                self.end_headers()
                self.wfile.write(data)
            else:
                _send_json(self, {"error": "Not found"}, 404)

        else:
            _send_json(self, {"error": "Not found"}, 404)

    # ── POST ──────────────────────────────────────────────────────────────────
    def do_POST(self):
        body = _read_body(self)

        if self.path == "/config":
            cfg = _load_cfg()
            cfg.update(body)
            _save_cfg(cfg)
            _send_json(self, {"ok": True})

        elif self.path == "/auth":
            token   = (body.get("token") or "").strip()
            correct = token and (_sha256(token) == _expected_token_hash())
            if correct:
                _save_auth({"token_hash": _sha256(token), "token_ok": True, "ts": time.time()})
            _send_json(self, {"ok": correct})

        elif self.path == "/model/validate":
            model_cfg = body
            _dep_err = _ensure_openai()
            if _dep_err:
                _send_json(self, {"ok": False, "error": _dep_err})
                return
            try:
                from istkc.models import make_client
                client = make_client(model_cfg)
                client.chat.completions.create(
                    model=model_cfg["api_id"],
                    messages=[{"role": "user", "content": "hi"}],
                    max_tokens=1,
                )
                _send_json(self, {"ok": True})
            except Exception as exc:
                _send_json(self, {"ok": False, "error": str(exc)})

        elif self.path == "/chat":
            self._stream_chat(body)

        elif self.path == "/run-brief":
            self._run_brief(body)

        elif self.path == "/brain/distil":
            self._distil(body)

        elif self.path == "/session/save":
            # Cheap archive — no LLM. Overwrites the conversation's own file so
            # one conversation stays one archive that updates in place.
            messages   = body.get("messages", [])
            session_id = body.get("session_id", "")
            if len(messages) < 1:
                _send_json(self, {"ok": False, "error": "no messages"})
            else:
                try:
                    with _lock:
                        p = _brain.save_session(messages, session_id=session_id)
                    _send_json(self, {"ok": True, "file": p.name})
                except Exception as e:
                    _send_json(self, {"ok": False, "error": str(e)})

        elif self.path == "/brain/clear":
            with _lock:
                _brain.clear()
            _send_json(self, {"ok": True})

        elif self.path == "/session/delete":
            from istkc.brain import SESSIONS_DIR
            fname = Path(body.get("file", "")).name   # basename only — no traversal
            p = SESSIONS_DIR / fname
            if fname and p.exists() and p.suffix == ".jsonl":
                try:
                    p.unlink()
                    _send_json(self, {"ok": True})
                except Exception as e:
                    _send_json(self, {"ok": False, "error": str(e)})
            else:
                _send_json(self, {"ok": False, "error": "not found"}, 404)

        elif self.path == "/workspace/index":
            workspace = body.get("workspace", "")
            if not workspace:
                _send_json(self, {"ok": False, "error": "workspace required"})
                return
            # Run index build in a background thread so the response is instant
            idx = get_or_create_index(workspace)
            def _build():
                idx.build()
            threading.Thread(target=_build, daemon=True).start()
            _send_json(self, {"ok": True, "message": "Indexing started in background"})

        elif self.path == "/task/run-agent":
            # Start a background agent task; return task ID immediately
            task = _new_task()
            model_cfg = body.get("model_cfg", {})
            brief     = body.get("brief", {})
            workspace = body.get("workspace", str(Path.home()))

            def _run():
                self._run_agent_task(task, brief, model_cfg, workspace)

            threading.Thread(target=_run, daemon=True).start()
            _send_json(self, {"ok": True, "task_id": task.id})

        else:
            _send_json(self, {"error": "Not found"}, 404)

    # ── SSE chat stream with tool-use loop ────────────────────────────────────
    def _stream_chat(self, body: dict) -> None:
        messages   = body.get("messages",  [])
        model_cfg  = body.get("model_cfg", {})
        workspace  = body.get("workspace", str(Path.home()))
        use_tools  = body.get("use_tools", True)

        self.send_response(200)
        _cors_headers(self)
        self.send_header("Content-Type",  "text/event-stream; charset=utf-8")
        self.send_header("Cache-Control", "no-cache")
        self.send_header("X-Accel-Buffering", "no")
        self.end_headers()

        def sse(data: dict) -> None:
            line = f"data: {json.dumps(data, ensure_ascii=False)}\n\n"
            self.wfile.write(line.encode("utf-8"))
            self.wfile.flush()

        try:
            _dep_err = _ensure_openai()
            if _dep_err:
                sse({"error": _dep_err})
                sse({"done": True})
                return
            from istkc.models import make_client
            client = make_client(model_cfg)

            provider = model_cfg.get("provider", "openai")
            api_id   = model_cfg.get("api_id", "")

            # Tool support: the renderer sends a registry-driven hint per model
            # (capsFor(api_id).tools). When present it wins — it knows about
            # tool-capable Ollama/ChatGPT models and tool-less reasoners alike.
            # Without a hint, fall back to the old provider heuristic.
            hint = body.get("supports_tools")
            if hint is not None:
                supports_tools = bool(use_tools and hint)
            else:
                supports_tools = (
                    use_tools
                    and provider in ("openai", "openrouter", "chatgpt")
                    and api_id not in ("o1", "o3-mini")
                )

            current_messages = list(messages)

            for _round in range(25):   # max 25 tool-call rounds (raised from 8 for complex multi-step workflows)
                kwargs: dict = dict(
                    model      = api_id,
                    messages   = current_messages,
                    stream     = True,
                    max_tokens = 8192,
                )
                if supports_tools:
                    kwargs["tools"]       = _get_tools()
                    kwargs["tool_choice"] = "auto"

                # Retry with exponential back-off on transient errors (#1)
                # Falls back to free models if primary exhausted (#6).
                # Fallback IDs are OpenRouter slugs — only usable on that provider;
                # sending them to OpenAI/Ollama would just 404 repeatedly.
                _FALLBACKS = [
                    "google/gemma-3-27b-it:free",
                    "mistralai/mistral-7b-instruct:free",
                ] if provider == "openrouter" else []
                _RETRY_DELAYS = [1, 3, 8]
                stream = None
                _last_exc: Exception | None = None
                for _attempt in range(len(_RETRY_DELAYS) + 1):
                    if _attempt:
                        time.sleep(_RETRY_DELAYS[_attempt - 1])
                    try:
                        stream = client.chat.completions.create(**kwargs)
                        break
                    except Exception as _e:
                        _last_exc = _e
                        _msg = str(_e).lower()
                        # Capability degradation — the endpoint can't do tools
                        # or images (common on :free OpenRouter endpoints).
                        # Strip the unsupported part and retry instead of
                        # failing the whole turn.
                        if "tools" in kwargs and "tool" in _msg and ("support" in _msg or "404" in _msg):
                            kwargs.pop("tools", None)
                            kwargs.pop("tool_choice", None)
                            supports_tools = False
                            sse({"text": "\n⚠️ This model's endpoint doesn't support tools — continuing without them.\n"})
                            continue
                        if "image input" in _msg or ("image" in _msg and "support" in _msg):
                            current_messages = _strip_image_parts(current_messages)
                            kwargs["messages"] = current_messages
                            sse({"text": "\n⚠️ This model can't see images — continuing with text only.\n"})
                            continue
                        _status = getattr(getattr(_e, "response", None), "status_code", None)
                        _transient = _status in (429, 500, 503) or "timeout" in str(_e).lower()
                        if not _transient:
                            raise
                if stream is None:
                    for _fb_model in _FALLBACKS:
                        try:
                            kwargs["model"] = _fb_model
                            stream = client.chat.completions.create(**kwargs)
                            sse({"text": f"\n⚠️ Primary model unavailable — switched to {_fb_model}\n"})
                            break
                        except Exception:
                            pass
                    if stream is None:
                        raise _last_exc  # type: ignore[misc]

                full_text:    str             = ""
                tool_calls_acc: dict[int, dict] = {}  # idx → {id, name, args_str}
                finish_reason: str | None     = None

                for chunk in stream:
                    if not chunk.choices:
                        continue
                    choice        = chunk.choices[0]
                    finish_reason = choice.finish_reason or finish_reason
                    delta         = choice.delta

                    # Text token
                    text_delta = (delta.content or "") if delta else ""
                    if text_delta:
                        full_text += text_delta
                        sse({"text": text_delta})

                    # Tool-call argument fragments
                    if supports_tools and delta and delta.tool_calls:
                        for tc in delta.tool_calls:
                            idx = tc.index
                            if idx not in tool_calls_acc:
                                tool_calls_acc[idx] = {"id": "", "name": "", "args_str": ""}
                            if tc.id:
                                tool_calls_acc[idx]["id"] += tc.id
                            if tc.function:
                                if tc.function.name:
                                    tool_calls_acc[idx]["name"] += tc.function.name
                                if tc.function.arguments:
                                    tool_calls_acc[idx]["args_str"] += tc.function.arguments

                # No tool calls → check whether the model paused mid-task before breaking
                if finish_reason != "tool_calls" or not tool_calls_acc:
                    # Detect prose that signals the model plans more work but forgot to call
                    # tools. Only the TAIL of the response counts — phrases like "step 2"
                    # or "now I'll" in the middle of a *finished* answer are normal prose,
                    # and nudging on them made NOVA ramble past a completed task.
                    _continuation_triggers = [
                        "now i'll", "now i need", "next i'll", "next i need",
                        "let me now", "i'll now", "i need to run", "i need to write",
                        "i'll run", "i'll write", "i'll create",
                        "i'll commit", "i'll add", "i'll test", "i'll save",
                        "now let's", "let's now",
                    ]
                    _tail = full_text.lower()[-300:]
                    _mid_task = bool(full_text) and any(t in _tail for t in _continuation_triggers)
                    if _mid_task and _round < 20:
                        # Model paused mid-task — nudge it to continue with tools
                        current_messages = current_messages + [
                            {"role": "assistant", "content": full_text or ""},
                            {"role": "user",      "content":
                             "Continue. Execute ALL remaining steps now using tool calls — "
                             "do not describe them, just do them one after another."},
                        ]
                        continue  # re-enter the loop
                    break

                # Execute each tool and gather results
                asst_tool_calls  = []
                tool_result_msgs = []

                for idx in sorted(tool_calls_acc.keys()):
                    tc   = tool_calls_acc[idx]
                    name = tc["name"]
                    try:
                        args = json.loads(tc["args_str"] or "{}")
                    except Exception:
                        args = {}

                    # Tell the frontend what NOVA is doing
                    sse({"tool_call": {"name": name, "args": args}})

                    # Store sse + model_cfg in thread-local so nested tools can use them
                    _tl.sse_cb    = sse
                    _tl.model_cfg = model_cfg
                    result = _execute_tool(name, args, workspace)

                    # Parse special browser event annotations from the result
                    # These trigger UI updates (screenshot display, preview pane navigation)
                    clean_lines = []
                    for line in result.split('\n'):
                        if line.startswith('SCREENSHOT_B64:'):
                            sse({"browser_screenshot": line[len('SCREENSHOT_B64:'):]})
                            clean_lines.append('📸 [screenshot]')
                        elif line.startswith('PREVIEW_URL:'):
                            sse({"preview_navigate": line[len('PREVIEW_URL:'):]})
                            clean_lines.append(line)  # keep in result for model to read
                        else:
                            clean_lines.append(line)
                    result = '\n'.join(clean_lines)

                    # Send a short preview back to the frontend
                    sse({"tool_result": {
                        "name":    name,
                        "snippet": result[:300] + ("…" if len(result) > 300 else ""),
                    }})

                    asst_tool_calls.append({
                        "id":       tc["id"],
                        "type":     "function",
                        "function": {"name": name, "arguments": tc["args_str"]},
                    })
                    tool_result_msgs.append({
                        "role":         "tool",
                        "tool_call_id": tc["id"],
                        "content":      result,
                    })

                # Append the assistant turn (with tool calls) + tool results before looping
                asst_msg: dict = {
                    "role":       "assistant",
                    "content":    full_text or None,
                    "tool_calls": asst_tool_calls,
                }
                current_messages = current_messages + [asst_msg] + tool_result_msgs

        except Exception as exc:
            sse({"error": str(exc)})

        sse({"done": True})

    # ── /run task brief ───────────────────────────────────────────────────────
    def _run_brief(self, body: dict) -> None:
        messages  = body.get("messages",  [])
        model_cfg = body.get("model_cfg", {})

        prompt = (
            "Based on this conversation, produce a precise task brief for the coding agent. "
            "Return ONLY a raw JSON object — no markdown fences, no prose:\n"
            '{"task":"one clear sentence","context":["bullet","..."],"files":["path/file.py"]}'
        )
        brief_msgs = (
            [m for m in messages if m["role"] == "system"]
            + [m for m in messages if m["role"] in ("user", "assistant")][-8:]
            + [{"role": "user", "content": prompt}]
        )

        try:
            from istkc.models import make_client
            client = make_client(model_cfg)
            resp   = client.chat.completions.create(
                model=model_cfg["api_id"],
                messages=brief_msgs,
                max_tokens=600,
                temperature=0.2,
            )
            raw = (resp.choices[0].message.content or "").strip()
            # Extract the outermost JSON object regardless of any surrounding
            # prose or markdown fences the model added.
            m = re.search(r"\{.*\}", raw, re.DOTALL)
            if m:
                raw = m.group(0)
            brief = json.loads(raw)
            _send_json(self, {"ok": True, "brief": brief})
        except Exception as exc:
            _send_json(self, {"ok": False, "error": str(exc)})

    # ── Background agent task runner ─────────────────────────────────────────
    def _run_agent_task(self, task: BgTask, brief: dict, model_cfg: dict, workspace: str) -> None:
        """
        Execute a multi-file coding task server-side.
        Emits progress events to `task.q` so the frontend can stream them.
        """
        from istkc.models import make_client
        files = brief.get("files", [])
        ctx   = brief.get("context", [])
        job   = brief.get("task", "Implement the requested changes")

        if not files:
            task.fail("No files specified in the task brief.")
            return

        try:
            client = make_client(model_cfg)
        except Exception as e:
            task.fail(str(e))
            return

        results: list[dict] = []
        ctx_str = "\n".join(f"- {c}" for c in ctx)
        _MAX_REWRITE_CHARS = 24_000   # beyond this a "full rewrite" would silently drop content

        for i, file_path in enumerate(files):
            fpath = _resolve(file_path, workspace)
            task.emit(step=i + 1, total=len(files), file=file_path, status="reading")

            current = ""
            try:
                current = fpath.read_text("utf-8", errors="replace")
            except FileNotFoundError:
                current = ""  # new file
            except Exception as e:
                task.emit(step=i + 1, total=len(files), file=file_path, status="error", error=str(e))
                continue

            # Guard: never regenerate a file the model can't fully see — the
            # unseen tail would be destroyed by the rewrite.
            if len(current) > _MAX_REWRITE_CHARS:
                err = (f"file is {len(current)} chars — too large for a safe full rewrite; "
                       f"skipped to avoid data loss (edit it in chat with edit_file instead)")
                results.append({"file": file_path, "ok": False, "error": err})
                task.emit(step=i + 1, total=len(files), file=file_path, status="error", error=err)
                continue

            prompt = (
                f"You are implementing this coding task:\n\n"
                f"**Task:** {job}\n\n"
                f"**Context:**\n{ctx_str}\n\n"
                f"**Current content of `{file_path}`:**\n"
                f"```\n{current}\n```\n\n"
                f"Write the COMPLETE new version of `{file_path}`. "
                f"Return ONLY the file content — no markdown fences, no explanation."
            )

            task.emit(step=i + 1, total=len(files), file=file_path, status="generating")

            new_content   = ""
            finish_reason = None
            try:
                stream = client.chat.completions.create(
                    model=model_cfg["api_id"],
                    messages=[{"role": "user", "content": prompt}],
                    stream=True,
                    max_tokens=8192,
                )
                for chunk in stream:
                    if not chunk.choices:
                        continue
                    choice        = chunk.choices[0]
                    finish_reason = choice.finish_reason or finish_reason
                    delta = (choice.delta.content or "") if choice.delta else ""
                    if delta:
                        new_content += delta
            except Exception as e:
                task.emit(step=i + 1, total=len(files), file=file_path, status="error", error=str(e))
                continue

            # Guard: a length-cut generation is an incomplete file — never write it.
            if finish_reason == "length":
                err = "model output hit the token limit — file NOT written (would be truncated)"
                results.append({"file": file_path, "ok": False, "error": err})
                task.emit(step=i + 1, total=len(files), file=file_path, status="error", error=err)
                continue

            new_content = _strip_md_fences(new_content)

            # Write the file
            task.emit(step=i + 1, total=len(files), file=file_path, status="writing")
            try:
                fpath.parent.mkdir(parents=True, exist_ok=True)
                fpath.write_text(new_content, encoding="utf-8")
                results.append({"file": file_path, "ok": True})
                task.emit(step=i + 1, total=len(files), file=file_path, status="done",
                          preview=new_content[:200])
            except Exception as e:
                results.append({"file": file_path, "ok": False, "error": str(e)})
                task.emit(step=i + 1, total=len(files), file=file_path, status="error", error=str(e))

        ok_count  = sum(1 for r in results if r.get("ok"))
        summary   = f"Agent complete: {ok_count}/{len(files)} files written successfully."
        failed    = [r for r in results if not r.get("ok")]
        if failed:
            summary += "\n" + "\n".join(f"✕ {r['file']}: {r.get('error', 'failed')}" for r in failed)
        task.finish(result=summary)

    # ── Brain distillation ────────────────────────────────────────────────────
    def _distil(self, body: dict) -> None:
        messages   = body.get("messages",  [])
        model_cfg  = body.get("model_cfg", {})
        session_id = body.get("session_id", "")
        try:
            from istkc.models import make_client
            client = make_client(model_cfg)
            with _lock:
                bd      = _brain.load()
                updated = _brain.distil(client, model_cfg["api_id"], messages, bd)
                _brain.save_session(messages, session_id=session_id)
                _brain.save(updated)
            _send_json(self, {"ok": True, "stats": _brain.stats(updated)})
        except Exception as exc:
            _send_json(self, {"ok": False, "error": str(exc)})


# ── Entry point ───────────────────────────────────────────────────────────────

# Load any skills that exist on disk at startup
_load_skills()

# Warm the core dependency in the background so the first chat doesn't stall
# on a pip install (macOS system Python ships without 'openai').
threading.Thread(target=_ensure_openai, daemon=True).start()


def run(port: int = PORT) -> None:
    try:
        server = ThreadingHTTPServer(("127.0.0.1", port), Handler)
    except OSError:
        # Port already in use — reuse ONLY if it's our bridge from the SAME app
        # version; a stale bridge from a previous install must not be adopted.
        try:
            import urllib.request as _ur, json as _jj
            resp = _ur.urlopen(f"http://127.0.0.1:{port}/health", timeout=2)
            data = _jj.loads(resp.read())
            if data.get("ok") and data.get("app_version") == _APP_VERSION:
                print(f"BRIDGE_READY:{port}", flush=True)
                return
        except Exception:
            pass
        print(f"[bridge] Port {port} already in use (or stale bridge). Exiting.", flush=True)
        sys.exit(1)
    print(f"BRIDGE_READY:{port}", flush=True)   # Electron reads this line
    server.serve_forever()


if __name__ == "__main__":
    run(int(sys.argv[1]) if len(sys.argv) > 1 else PORT)
